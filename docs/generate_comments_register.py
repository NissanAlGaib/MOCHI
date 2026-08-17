"""Generate the MOCHI Comments & Action Register as a Word document.

    python docs/generate_comments_register.py

Regenerate whenever status changes so the document never drifts from reality.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent / "MOCHI_Comments_Register.docx"

STATUS_COLOURS = {
    "IMPLEMENTED": RGBColor(0x1B, 0x6E, 0x1B),
    "PARTIAL": RGBColor(0xB0, 0x6A, 0x00),
    "PLANNED": RGBColor(0x1F, 0x4E, 0x9C),
    "NOT STARTED": RGBColor(0x88, 0x88, 0x88),
    "DECISION NEEDED": RGBColor(0xB0, 0x1C, 0x1C),
    "THESIS EDIT": RGBColor(0x6A, 0x1B, 0x9A),
}

# ---------------------------------------------------------------------------
# Register content: (ID, Comment / Issue, Source, Action / How addressed,
#                    Status, Evidence)
# ---------------------------------------------------------------------------

ADVISER = [
    ("A1", "Use NLP for parsing strings and identifying string parts",
     "Adviser keyword: NLP",
     "Adopt spaCy for parsing. Proposed as a new Stage 1.5 that detects "
     "imperative verb + instruction-object structures generically, rather than "
     "matching literal phrasings. Directly targets the measured Stage I recall "
     "gap of 5.2%.",
     "PLANNED", "Proposed; not yet built. Would follow Phase 6."),

    ("A2", "Remove parts of speech / POS tagging",
     "Adviser keyword: Remove parts of speech, post-tagging",
     "Split into two uses. (a) POS tagging AS A DETECTION SIGNAL - adopted, see "
     "A1. (b) POS/stopword REMOVAL to shrink prompts - NOT adopted. Injection "
     "lives in imperative verb + object structure; removing function words "
     "destroys the signal and creates a train/serve mismatch because runtime "
     "text is never stripped. MOCHI must also forward the original prompt "
     "unaltered or it breaks its own transparency claim.",
     "DECISION NEEDED", "Rationale documented; needs adviser sign-off on (b)."),

    ("A3", "Use NLTK",
     "Adviser keyword: NLTK",
     "Conflict with the thesis: Appendix C already specifies spaCy as the "
     "Heuristic Parsing Engine. Recommendation is spaCy for runtime parsing "
     "(1-3 ms/doc vs 5-20 ms for NLTK, matters for the NFR1 budget), with NLTK "
     "retained only for WordNet-based augmentation. Appendix C or the adviser "
     "note must be reconciled.",
     "DECISION NEEDED", "Appendix C, System Configuration table."),

    ("A4", "Augmentation algorithms",
     "Adviser keyword: Augmentation algo",
     "Plan: nlpaug (character/word/sentence level) plus EDA (synonym "
     "replacement, random insertion/swap/deletion) as the citable method. "
     "Character-level augmentation generates exactly the evasions the Phase 3 "
     "normalization layer defeats, enabling a with/without ablation. "
     "CONSTRAINT: augment the TRAINING split only, after deduplication and "
     "splitting, or variants leak across splits and inflate Stage II results.",
     "PLANNED", "For Phase 8 training data preparation."),

    ("A5", "Use an LLM to generate mimicked prompt injections",
     "Adviser keyword: Use llm to generate mimiced prompt injections",
     "Two-part approach. (a) Prefer PARAPHRASING existing benchmark attacks "
     "over asking a model to invent them - frontier models frequently refuse, "
     "which silently biases the corpus toward weak samples. (b) For generation, "
     "use a local open-weight model via Ollama (mistral:7b-instruct, "
     "qwen2.5:7b-instruct, or llama3.1:8b-instruct) - no refusals, no API cost, "
     "reproducible offline. The gateway already speaks the OpenAI API, so "
     "OPENAI_BASE_URL can point at Ollama and reuse the whole pipeline. "
     "CONSTRAINTS: generated attacks must be validated against an undefended "
     "model before being labelled malicious, and must never supply headline "
     "detection metrics - published benchmarks do that.",
     "PLANNED", "For Phase 13 attack simulation and Phase 8 augmentation."),

    ("A6", "Duplicate detection within prompts (\"n-max / n-pax\")",
     "Adviser keyword, term uncertain",
     "Interpreted as N-GRAMS. Relevant algorithm family: n-grams for "
     "tokenisation, MinHash + LSH (datasketch) or SimHash for near-duplicate "
     "detection at corpus scale, Jaccard similarity for overlap. This mirrors "
     "the 'Semantic Deduplication' step in MCP-Guard's filtration pipeline. "
     "A normalized-hash near-duplicate check is already implemented; MinHash "
     "would be the scale-up if needed.",
     "PARTIAL", "eval/audit_datasets.py, eval/clean_datasets.py"),

    ("A7", "Token management / reduce token load",
     "Adviser inquiry",
     "MEASURED, and the premise does not hold for this corpus. Artifact "
     "cleaning yields only 0.6% token reduction (31.49M -> 31.30M), and only "
     "~1% of samples exceed the 512-token Stage II window. The real efficiency "
     "lever is corpus selection (subsampling jayavibhav), not text compression. "
     "Legitimate token management remains: the 4,096-token inspection window "
     "with chunking (already in Table 7), and compressing Stage III arbiter "
     "input only (LLMLingua) in Phase 9.",
     "IMPLEMENTED", "reports/audit.json; eval/clean_datasets.py --dry-run"),

    ("A8", "Target LLM is a moving target - MOCHI risks obsolescence",
     "Adviser objection",
     "Answered architecturally: MOCHI treats the target LLM as a black box, "
     "never touching weights or internals, so the protected model can change "
     "freely. A component-by-component dependence table was produced showing "
     "only the Stage III arbiter is version-pinned (one config value). "
     "Recommended additions: a 'Model Independence and Longevity' subsection "
     "under Validity, a retraining cadence in Recommendations, and a held-out "
     "4th LLM generalization test in Phase 4 evaluation.",
     "PARTIAL", "docs/ARCHITECTURE.md model-dependence table; thesis text pending."),
]

DATA = [
    ("D1", "MCP-ATTACKBENCH is not publicly available",
     "Verification of thesis Table 11",
     "Confirmed by checking the arXiv abstract, full paper text, GitHub and "
     "Hugging Face: the MCP-Guard paper states the dataset 'will be released' "
     "with no link or date. It accounted for 70,448 of the planned 97,448 "
     "samples (72%) and backed the Stage II training set. Substituted with "
     "PromptShield, deepset/prompt-injections and jayavibhav/prompt-injection.",
     "IMPLEMENTED", "Datasets downloaded to data/; Table 11 needs revision (T1)."),

    ("D2", "deepset/prompt-injections is entirely contained in jayavibhav",
     "Dataset audit",
     "RESOLVED - dropped. All 546 deepset texts appear verbatim inside "
     "jayavibhav (100% overlap, SHA-1 verified), so it contributes zero new "
     "examples and listing it as an independent source would inflate the "
     "reported corpus size. Removed from eval/fetch_datasets.py SOURCES and "
     "recorded in an EXCLUDED table with the reason; the file itself is moved "
     "to data/excluded/ rather than deleted so the claim stays re-verifiable.",
     "IMPLEMENTED", "eval/fetch_datasets.py EXCLUDED; data/excluded/deepset.csv."),

    ("D3", "Cross-dataset overlap creates train/test leakage",
     "Dataset audit",
     "Measured: 17 exact and 97 near duplicates between jayavibhav and "
     "promptshield_test; 29 near duplicates between promptshield_train and "
     "promptshield_validation. Training on jayavibhav and testing on "
     "promptshield_test would leak ~114 samples. Deduplication before "
     "splitting is now applied to the corpus: 358 exact, 420 near, 319 "
     "label-conflict and 76 empty rows removed (305,162 -> 303,989).",
     "IMPLEMENTED", "data/clean/ written by eval/clean_datasets.py."),

    ("D4", "254-322 rows carry contradictory labels",
     "Dataset audit",
     "Same normalized text appears with both benign and malicious labels. "
     "Cleaner drops these by default (--keep-conflicts to retain). Note the "
     "near-duplicate key strips punctuation, so this is an upper bound.",
     "IMPLEMENTED", "eval/clean_datasets.py, drop_conflicts=True"),

    ("D5", "Head-truncation would discard ~15% of attacks",
     "Dataset audit, signal-position analysis",
     "Of long malicious samples where Stage I located the payload: 54.5% head, "
     "30.6% middle, 14.8% tail. HuggingFace's default truncation=True keeps the "
     "FIRST N tokens and would silently drop tail-appended injections - exactly "
     "the Table 16 Direct Attack pattern. Must use the chunk-and-take-max "
     "approach already specified in Table 7. Caveat: based on 310 locatable "
     "samples of 4,001 sampled, so indicative rather than precise.",
     "PLANNED", "Constraint recorded for Phase 8 implementation."),

    ("D6", "jayavibhav dominates the corpus and is not application-representative",
     "Dataset audit",
     "261,737 of 305,708 samples (86%). Its benign half is general "
     "instruction-following text (maths word problems etc.), not the banking / "
     "travel / workspace / email traffic the thesis describes. RESOLVED - "
     "capped at 40,000 rows after deduplication, stratified by label and "
     "seeded (seed 42), giving 40,000 jayavibhav against 42,765 PromptShield: "
     "rough parity, so neither source dictates the Stage II decision boundary. "
     "The cap preserves jayavibhav's own class ratio rather than forcing 50/50, "
     "so it remains a size decision only. Headline FPR claims should still be "
     "weighted toward PromptShield, whose benign samples are closer to "
     "realistic application traffic.",
     "IMPLEMENTED", "DATASET_CAPS in eval/clean_datasets.py; --no-cap disables."),

    ("D7", "PromptShield identity on Hugging Face was ambiguous",
     "Dataset acquisition",
     "Four datasets named PromptShield exist on the Hub. Selected "
     "hendzh/PromptShield: oldest upload, 551 downloads vs 24-72 for the "
     "others, and its README links arXiv 2501.15145 - the paper cited as [22]. "
     "The other three are later mirrors with byte-identical READMEs. Cite the "
     "arXiv preprint ID alongside the ACM CODASPY DOI.",
     "IMPLEMENTED", "eval/fetch_datasets.py SOURCES, with provenance comment."),

    ("D12", "Span-level redaction left the exfiltration target in the prompt",
     "Found by running demo/enforcement_demo.py",
     "DEFECT, now FIXED. Sanitisation originally removed the detector's matched "
     "span. But a match is a partial view of the instruction: the pattern for "
     "'email X to Y' matched 'email the admin password' and not "
     "'to evil@example.com', so redaction forwarded the exfiltration address to "
     "the model, plus fragments stranded between two overlapping matches. "
     "Redaction now removes the enclosing sentence, and the boundary rule "
     "requires whitespace after a terminator so the dot in 'example.com' does "
     "not end a sentence mid-domain. Lesson for the write-up: removing what a "
     "detector matched is not the same as neutralising a payload.",
     "IMPLEMENTED", "mochi/mitigate/sanitizer.py; 6 regression tests."),

    ("D9", "Stage I silently truncated long documents at 20,000 characters",
     "Adviser comment: 'what if 10 words but 1 word only malicious'",
     "DEFECT, now FIXED. scan_text() cut every text to the first 20,000 chars "
     "and still reported a clean 'pass', so any injection past that point was "
     "invisible - the same tail-truncation failure D5 flagged for Stage II, "
     "one stage earlier than anyone was looking. Its docstring claimed 'longer "
     "input is chunked by the caller'; no caller chunked. Now scanned as "
     "overlapping 20,000-char windows (512-char overlap so a boundary-straddling "
     "payload is still seen whole), bounded at 100,000 chars total, and a scan "
     "that hits the ceiling reports 'pass_truncated' instead of 'pass'. "
     "Measured cost ~0.5 ms per 1,000 chars, so the earlier '<2 ms budget' "
     "claim held only for short prompts.",
     "IMPLEMENTED", "mochi/detect/chunking.py; 11 regression tests in "
                    "tests/test_stage1.py."),

    ("D10", "Malicious spans occupy a median 3.4% of their document",
     "Adviser comments: 'context to the class label', 'which token can "
     "contribute to the prediction'",
     "MEASURED. Using Stage I's matched span as a proxy across 2,624 locatable "
     "malicious samples: 72.0% of attacks occupy under 5% of their document, "
     "21.0% occupy 5-10%, median 3.4%, worst case 0.0006%. The adviser's "
     "'1 word in 10' is generous - reality is nearer 1 in 30. Consequence: "
     "mean pooling (the sentence-transformers default) yields a document vector "
     "that is ~97% benign text, so the malicious direction is averaged into "
     "noise. Stage II therefore uses gated attention pooling and max-over-"
     "windows, and reports the winning span plus top attributed tokens. "
     "Formal framing for the thesis: multiple-instance learning (Ilse, Tomczak "
     "& Welling, ICML 2018) - the bag carries the label, a few instances carry "
     "the evidence. Caveat: the 2,624 sample is biased toward attacks Stage I "
     "already catches; re-measure once Stage II can score spans on all 35,838.",
     "IMPLEMENTED", "training/model.py AttentionPool; "
                    "mochi/detect/stage2_semantic.py."),

    ("D11", "The corpus cannot exercise the dilution problem it revealed",
     "Follows from D9/D10",
     "Only 3 of 82,765 cleaned samples exceed 20,000 characters, and 1 of "
     "those is malicious. The corpus is almost entirely short prompts, so the "
     "long-document truncation fix changes no corpus metric - and the "
     "dilution behaviour that matters most for indirect injection via retrieved "
     "documents cannot be measured on it at all. Report dilution handling from "
     "the synthetic tests, and state plainly in Chapter III that the public "
     "corpora are not representative of the indirect-injection deployment "
     "scenario MOCHI targets. This is an argument for a purpose-built indirect "
     "fixture set, which is what the Figure 12 attack simulation should become.",
     "DECISION NEEDED", "Affects Chapter III Limitations and Phase 13 design."),

    ("D8", "PromptShield ships official train/validation/test splits",
     "Dataset acquisition",
     "18,909 / 1,000 / 23,516. Overlap between its own train and test is 1 row - "
     "the split is clean. Using the official splits makes results directly "
     "comparable to published PromptShield numbers, which is stronger than "
     "re-splitting. Note promptshield_test is deliberately imbalanced "
     "(72% benign) to model realistic traffic - do not rebalance it.",
     "IMPLEMENTED", "data/promptshield_{train,validation,test}.csv"),
]

THESIS = [
    ("T1", "Table 11 Dataset Composition cites unavailable data",
     "Follows from D1",
     "Replace the MCP-ATTACKBENCH / PromptShield / AgentDojo rows and the "
     "97,448 total with the final corpus as built: PromptShield 42,765 "
     "(official splits, post-dedup) and jayavibhav 40,000 (capped, post-dedup), "
     "total 82,765 - 46,927 benign / 35,838 malicious, 43.3% malicious. "
     "State that MCP-ATTACKBENCH was unavailable at time of writing, and that "
     "deepset/prompt-injections was excluded as a proper subset of jayavibhav.",
     "THESIS EDIT", "Chapter III, Table 11. Numbers from data/clean/."),

    ("T2", "Table 12 Data Splits numbers are placeholders",
     "Follows from D8",
     "Replace 60,000 / 10,000 / 27,448 with PromptShield's actual official "
     "split sizes, or state the constructed split explicitly if combining "
     "corpora.",
     "THESIS EDIT", "Chapter III, Table 12."),

    ("T3", "Chapter III has no preprocessing / normalization stage",
     "Implementation finding",
     "The pipeline goes straight from interception to Stage I, but without "
     "normalization both Stage I and Stage II are blind to base64, zero-width "
     "splitting, homoglyphs and hidden HTML. Measured: a conventional pipeline "
     "misses 10 of 13 documented evasions; with normalization, 0 of 13. Needs a "
     "'Preprocessing and Normalization' subsection before Stage I, and "
     "normalization_flags added to the Figure 9 log schema.",
     "THESIS EDIT", "Chapter III, System Architecture and Algorithms."),

    ("T4", "Source tag taxonomy is missing a sixth tag",
     "Implementation finding",
     "Prior assistant output is replayed into context on multi-turn "
     "conversations, so an injection that succeeded on turn 1 re-enters on turn "
     "2. Added assistant_output (untrusted). Directly relevant to the "
     "Multi-Step Attack Chain row in Table 16.",
     "IMPLEMENTED", "mochi/detect/segments.py SourceTag; thesis text pending."),

    ("T5", "'Continuous, stateless per-request' contradicts multi-step detection",
     "Implementation finding",
     "Table 16 includes multi-step attack chains, which a stateless design "
     "cannot detect by construction - each turn is scored in isolation. "
     "Agreed resolution: add a session-level rolling risk accumulator (Phase 7) "
     "and revise the text to say per-request DETECTION is stateless while "
     "session-level ESCALATION is intentionally stateful.",
     "PLANNED", "Phase 7; approach agreed, not yet built."),

    ("T6", "Several Table 8 patterns would destroy the FPR target",
     "Phase 6 implementation",
     "As literally written: (--|\\bOR\\b) matches the English word 'or'; "
     "\\bAPI_KEY\\b matches ordinary developer questions; \\bDAN\\b matches the "
     "name Dan; \\b(sh|bash|rm)\\b matches shell tutorials. Implemented with "
     "severity tiers (only 'high' blocks) and tightened context. Measured "
     "result: FPR 0.0015 against a <0.01 target. Table 8 should be updated to "
     "the implemented patterns, or a note added explaining the refinement.",
     "IMPLEMENTED", "mochi/patterns.json; tests/test_stage1.py near-miss cases."),

    ("T7", "NFR3 (>95% F1) reads as a Stage I requirement",
     "Phase 6 measurement",
     "Stage I alone measures F1 = 0.0989 (precision 0.97, recall 0.052). This "
     "is the expected and correct behaviour of a rule-based filter and matches "
     "the limitation the literature review already attributes to rule-based "
     "detection. State explicitly that NFR3 applies to full MOCHI, not to "
     "individual stages, so the ablation row is not read as a failed "
     "requirement.",
     "THESIS EDIT", "Chapter III, Table 6 (Non-Functional Requirements)."),

    ("T8", "ASR and Mitigation Rate are complements, not independent findings",
     "Phase 5 implementation",
     "Measured from one confusion matrix they sum to exactly 1.0; there is a "
     "test asserting this. The thesis lists them as two primary metrics under "
     "separate hypotheses (H05, H06). Add a sentence acknowledging the "
     "relationship, or a panelist will raise it.",
     "THESIS EDIT", "Chapter III, Table 17 and the H05/H06 statements."),

    ("T9", "Normality is asserted, not tested",
     "Phase 5 implementation",
     "The paired t-test is justified partly on data being 'essentially normally "
     "distributed'. A Shapiro-Wilk check on the paired differences now runs "
     "automatically and warns on failure; wilcoxon_test() provides the "
     "non-parametric fallback. Report the check rather than asserting the "
     "assumption. Also note n=5 with 5-fold CV, and that very large Cohen's d "
     "values are an artifact of a null control.",
     "IMPLEMENTED", "eval/stats.py check_normality(), wilcoxon_test()"),

    ("T10", "Stage III latency is unbudgeted",
     "Architecture review",
     "NFR1 budgets <2 ms syntactic and <55 ms semantic but says nothing about "
     "Stage III, which makes a network call (300-800 ms) and cannot meet either. "
     "Latency is already recorded per stage. Add an explicit statement that "
     "Stage III fires only in the uncertain band, and report the escalation "
     "rate alongside it.",
     "PARTIAL", "Timing infrastructure exists (stage_timer); text pending."),

    ("T11", "Figure 8 log path does not match implementation",
     "Phase 2 implementation",
     "Figure 8 specifies logs/mochi.log; the implementation writes JSON Lines to "
     "logs/mochi.jsonl, which is the accurate extension for the format and "
     "matters for tooling. Cosmetic - update Figure 8 or leave as configurable.",
     "THESIS EDIT", "Chapter III, Figure 8."),

    ("T12", "Source tagging appears to contradict the zero-code-change claim",
     "Architecture review",
     "The thesis states MOCHI 'necessitates' structured JSON with metadata "
     "tagging, which IS an application change and conflicts with NFR2. "
     "Resolved in implementation: tagging is OPTIONAL. Untagged requests fall "
     "back to role-based inference (system/user/tool), so integration remains a "
     "one-line base_url change. Text should present tagging as an enhanced mode.",
     "IMPLEMENTED", "mochi/detect/segments.py segment_messages(); thesis pending."),
]

PHASES = [
    ("0", "Environment setup", "IMPLEMENTED", "venv, requirements, repo layout"),
    ("1", "Gateway skeleton (OpenAI-compatible proxy)", "IMPLEMENTED",
     "Verified end-to-end against the live OpenAI API"),
    ("2", "Structured JSON telemetry", "IMPLEMENTED",
     "logs/mochi.jsonl; payload hashing on by default"),
    ("3", "Normalization / de-obfuscation", "IMPLEMENTED",
     "13 evasion techniques defeated, 0 missed; 2 benign controls clean"),
    ("4", "Source tagging and payload parsing", "IMPLEMENTED",
     "Trust tiers, direct/indirect classification"),
    ("5", "Evaluation harness", "IMPLEMENTED",
     "Metrics, 5-fold CV, H01-H06 significance testing"),
    ("6", "Stage I syntactic filtering", "IMPLEMENTED",
     "On the balanced corpus: P=0.9728 R=0.0590 F1=0.1112 FPR=0.0013, "
     "0.550 ms mean / 1.611 ms p95. Long-document scanning fixed (see D9)."),
    ("7", "Session risk accumulator", "NOT STARTED",
     "Approach agreed; addresses T5 and multi-step chains"),
    ("8", "Stage II semantic detection", "PARTIAL",
     "Code complete and fully tested through a stub scorer (44 tests); model "
     "not yet trained. Chunk-and-take-max, gated attention pooling, and span "
     "attribution implemented per D5/D10. Remaining: Colab fine-tune and "
     "export to models/e5-fine-tuned/."),
    ("9", "Stage III cognitive arbitration", "NOT STARTED", "LLM arbiter"),
    ("10", "Enforcement and sanitization", "IMPLEMENTED",
     "ALLOW / BLOCK / SANITIZE using segment trust. Closes FR2 - before this "
     "the gateway detected attacks and forwarded them anyway. Resolves the "
     "Stage II uncertain band by trust rather than an LLM arbiter (Q7). "
     "31 tests; demo/enforcement_demo.py covers 9 scenarios."),
    ("11", "Outbound interception", "NOT STARTED",
     "Leaked-instruction and URL-exfiltration scanning"),
    ("12", "Multi-provider adapters", "PARTIAL",
     "Interface built; only the OpenAI adapter exists"),
    ("13", "Full evaluation", "NOT STARTED", "Chapters IV tables"),
    ("14", "Packaging", "NOT STARTED", "Docker, User's Manual"),
]

DECISIONS = [
    ("Q1", "spaCy or NLTK?",
     "Appendix C says spaCy; adviser said NLTK. Recommendation: spaCy for "
     "runtime parsing, NLTK only for WordNet augmentation."),
    ("Q2", "Is POS/stopword removal from prompts approved or rejected?",
     "Recommendation: reject for prompt content (destroys detection signal and "
     "breaks transparency); accept POS tagging as a detection signal instead."),
    ("Q3", "Keep or drop deepset? - ANSWERED: drop",
     "Dropped. See D2. Retain the citation in Chapter II with the subset "
     "relationship to jayavibhav disclosed."),
    ("Q4", "Subsample jayavibhav? - ANSWERED: yes, capped at 40,000",
     "Capped. See D6. Final corpus 82,765 at 43.3% malicious."),
    ("Q7", "Omit the Stage III LLM arbiter? - ANSWERED: yes, by default",
     "Implemented in Phase 10: the Stage II uncertain band is resolved by "
     "source trust, deterministically and at zero latency. Stage III remains "
     "planned behind a flag, default off, reported as an ablation arm. Reasons "
     "beyond cost: an LLM arbiter is itself prompt-injectable, is not "
     "reproducible (which the Reliability section requires), and forces a "
     "network dependency on a security component."),
    ("Q8", "Build a purpose-built indirect-injection fixture set? - see D11",
     "The public corpora are 99.996% short prompts, so they cannot test the "
     "long-document dilution case that is MOCHI's main claimed value. "
     "Recommendation: yes - fold it into the Figure 12 attack simulation."),
    ("Q5", "Build the spaCy imperative detector (Stage 1.5)?",
     "Highest-value remaining lever on the 5.2% Stage I recall, and a "
     "defensible research contribution distinct from more regex."),
    ("Q6", "Which local LLM for generating mimicked injections?",
     "Recommendation: Ollama with mistral:7b-instruct or qwen2.5:7b-instruct."),
]


# ---------------------------------------------------------------------------
# Document construction
# ---------------------------------------------------------------------------


def add_status_run(cell, status: str) -> None:
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(status)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = STATUS_COLOURS.get(status, RGBColor(0, 0, 0))


def add_register_table(document: Document, rows, headers, widths) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for index, text in enumerate(headers):
        run = header_cells[index].paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(9)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            if headers[index] == "Status":
                add_status_run(cells[index], value)
                continue
            run = cells[index].paragraphs[0].add_run(str(value))
            run.font.size = Pt(8.5)
            if index == 0:
                run.bold = True

    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)


def heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def build() -> None:
    document = Document()

    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Inches(0.6)
    section.top_margin = section.bottom_margin = Inches(0.6)

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # --- title ---
    title = document.add_heading("MOCHI - Comments and Action Register", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "A Lightweight Deployable Middleware for Detecting and Mitigating "
        "Prompt Injection Attacks in Large Language Model Applications\n"
        "Lao, H. A. A.  |  Macias, J. M. C.  |  Maglangit, R. M.\n"
        f"Generated {date.today():%d %B %Y}"
    )
    run.italic = True
    run.font.size = Pt(10)

    document.add_paragraph()

    # --- legend ---
    heading(document, "How to read this document", level=1)
    legend = document.add_paragraph()
    legend.add_run("Status values:\n").bold = True
    for status, meaning in [
        ("IMPLEMENTED", "built, tested, and verified in the codebase"),
        ("PARTIAL", "partly built, or built in code but not yet reflected in the thesis"),
        ("PLANNED", "approach agreed and documented, implementation scheduled"),
        ("NOT STARTED", "no work done yet"),
        ("DECISION NEEDED", "blocked on a choice by the researchers or adviser"),
        ("THESIS EDIT", "no code change required; the manuscript needs revising"),
    ]:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(status)
        run.bold = True
        run.font.color.rgb = STATUS_COLOURS[status]
        paragraph.add_run(f" - {meaning}")

    note = document.add_paragraph()
    note.add_run(
        "All measured figures in this register come from runs against the "
        "assembled corpus of 305,708 labelled samples. Regenerate this document "
        "with docs/generate_comments_register.py after any status change."
    ).italic = True

    headers = ["ID", "Comment / Issue", "Source", "Action / How addressed",
               "Status", "Evidence"]
    widths = [0.4, 1.9, 1.1, 4.1, 1.0, 1.7]

    document.add_page_break()
    heading(document, "Section A - Adviser Comments", level=1)
    document.add_paragraph(
        "Comments raised by the adviser, with the action taken or proposed."
    )
    add_register_table(document, ADVISER, headers, widths)

    document.add_page_break()
    heading(document, "Section B - Dataset and Data-Quality Findings", level=1)
    document.add_paragraph(
        "Findings from verification of the datasets named in the methodology and "
        "from auditing the corpus actually assembled."
    )
    add_register_table(document, DATA, headers, widths)

    document.add_page_break()
    heading(document, "Section C - Thesis Manuscript Revisions", level=1)
    document.add_paragraph(
        "Discrepancies between the written methodology and the implemented "
        "system, or claims that measurement has since qualified. These require "
        "edits to the manuscript rather than to the code."
    )
    add_register_table(document, THESIS, headers, widths)

    document.add_page_break()
    heading(document, "Section D - Implementation Status by Phase", level=1)
    document.add_paragraph(
        "Development is organised into 14 phases (see docs/BUILD_PLAN.md). "
        "Phases 0-6 are complete and verified by 193 automated tests."
    )
    add_register_table(
        document, PHASES,
        ["Phase", "Description", "Status", "Notes / measured result"],
        [0.6, 3.0, 1.2, 5.4],
    )

    document.add_page_break()
    heading(document, "Section E - Open Decisions", level=1)
    document.add_paragraph(
        "Items requiring a decision from the researchers or adviser before the "
        "affected work can proceed."
    )
    add_register_table(
        document, DECISIONS,
        ["ID", "Question", "Recommendation"],
        [0.5, 3.2, 6.5],
    )

    document.add_page_break()
    heading(document, "Section F - Measured Results to Date", level=1)
    document.add_paragraph(
        "Figures available for Chapter IV as of this revision. Detection "
        "effectiveness for Stages II and III is not yet measurable."
    )

    results = [
        ("Corpus assembled", "305,708 samples (159,832 benign / 145,876 malicious)"),
        ("Baseline (no defence) accuracy", "0.5065"),
        ("Baseline Attack Success Rate", "1.0000 (by construction)"),
        ("Stage I accuracy", "0.5469"),
        ("Stage I precision", "0.9700"),
        ("Stage I recall", "0.0521"),
        ("Stage I F1", "0.0989"),
        ("Stage I false positive rate", "0.0015  (target < 0.0100 - MET)"),
        ("Stage I latency", "0.430 ms mean, 0.883 ms p95  (target < 2 ms - MET)"),
        ("Obfuscation resistance", "13 of 13 evasions revealed; conventional "
                                   "pipeline misses 10 of 13"),
        ("Automated tests", "193 passing"),
        ("Code coverage", "87%"),
    ]
    add_register_table(document, results, ["Measure", "Value"], [3.5, 6.7])

    interpretation = document.add_paragraph()
    interpretation.add_run("Interpretation of the Stage I recall figure: ").bold = True
    interpretation.add_run(
        "5.2% recall is the expected and correct behaviour of a rule-based "
        "filter, not a defect. Regex can only match phrasings already written "
        "down; the 138,275 missed attacks are paraphrases and semantic variants. "
        "This is precisely the limitation the Review of Related Literature "
        "attributes to rule-based detection - now measured on this system rather "
        "than cited. It is the empirical justification for the cascaded "
        "architecture and for Stage II."
    )

    document.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    print(f"  Section A - adviser comments      : {len(ADVISER)}")
    print(f"  Section B - dataset findings      : {len(DATA)}")
    print(f"  Section C - thesis revisions      : {len(THESIS)}")
    print(f"  Section D - phases                : {len(PHASES)}")
    print(f"  Section E - open decisions        : {len(DECISIONS)}")


if __name__ == "__main__":
    build()
